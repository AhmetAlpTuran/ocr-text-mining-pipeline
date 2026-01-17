import argparse
from pathlib import Path

import pandas as pd
from docx import Document


def build_docx_table(df: pd.DataFrame, title: str) -> Document:
    document = Document()
    if title:
        document.add_heading(title, level=1)

    if df.empty:
        document.add_paragraph("No data available.")
        return document

    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, col_name in enumerate(df.columns):
        header_cells[idx].text = str(col_name)

    for row in df.itertuples(index=False):
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            row_cells[idx].text = "" if pd.isna(value) else str(value)

    return document


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Export summary CSV to a DOCX table."
    )
    parser.add_argument(
        "--input",
        default=None,
        help="Path to summary CSV (ozet).",
    )
    parser.add_argument(
        "--output",
        default=None,
        help="Output DOCX path.",
    )
    parser.add_argument(
        "--title",
        default="Summary",
        help="Optional title for the document.",
    )
    args = parser.parse_args()

    if not args.input:
        raise ValueError("Please provide --input for the summary CSV.")

    input_path = Path(args.input)
    output_path = Path(args.output) if args.output else input_path.with_suffix(".docx")

    df = pd.read_csv(input_path)
    document = build_docx_table(df, args.title)
    document.save(output_path)
    print(f"Saved DOCX to {output_path}")


if __name__ == "__main__":
    main()
