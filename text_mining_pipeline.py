import argparse
import re
from pathlib import Path

import nltk
import ollama
import pandas as pd
import pdfplumber
from docx import Document
from nltk.tokenize import sent_tokenize
from tqdm import tqdm


TURKISH_LOWER_MAP = str.maketrans({"I": "ı", "İ": "i"})

CONTEXTUAL_RULES = [
    {"term": "Burjuva", "rule_type": "non_european"},
    {"term": "Burjuvazi", "rule_type": "non_european"},
    {"term": "Constantinopolis", "rule_type": "post_1453"},
    {"term": "Feodalizm", "rule_type": "non_european"},
    {"term": "Feodalite", "rule_type": "non_european"},
    {"term": "Megaron", "rule_type": "megaron_non_greek"},
    {"term": "Mongoloid", "rule_type": "derogatory"},
    {"term": "Reform", "rule_type": "non_european"},
    {"term": "Rönesans", "rule_type": "non_european"},
    {"term": "Şövalye", "rule_type": "non_european"},
    {"term": "Şövalye Kültürü", "rule_type": "non_european"},
    {"term": "Siyahi", "rule_type": "derogatory"},
    {"term": "Zenci", "rule_type": "derogatory"},
    {"term": "Sarazen", "rule_type": "sarazen"},
]

NON_EUROPEAN_MARKERS = [
    "türk",
    "türkiye",
    "osmanlı",
    "arap",
    "iran",
    "fars",
    "hint",
    "hindistan",
    "çin",
    "japon",
    "afrika",
    "asya",
    "orta doğu",
    "ortadoğu",
    "orta asya",
    "kafkas",
    "kafkasya",
    "kürt",
    "pakistan",
    "afgan",
    "moğol",
]

DEROGATORY_MARKERS = [
    "aşağı",
    "aşağılık",
    "ilkel",
    "vahşi",
    "barbar",
    "geri",
    "yabani",
    "ötekileştir",
]

POST_1453_MARKERS = [
    "1453",
    "osmanlı",
    "istanbul",
    "fetih",
    "fatih",
]

MEGARON_NON_GREEK_MARKERS = [
    "sümer",
    "sümerler",
    "hitit",
    "hititler",
    "frig",
    "frigler",
    "frik",
    "frikler",
]

SARAZEN_CRUSADER_MARKERS = ["haçlı", "haçlılar", "hacli", "haclılar"]
SARAZEN_MUSLIM_MARKERS = ["müslüman", "musluman"]
SARAZEN_LABELING_MARKERS = ["tanımla", "adlandır", "isimlendir"]

CATEGORY_ORDER = [
    "Çağ ayrımı",
    "Gregoryen Miladı",
    "Roma Rakamları",
    "Olgusal Kavramlar",
    "Bağlamsal Kavramlar",
]
ROMAN_NUMERAL_PATTERN = re.compile(r"^[IVXLCDM]+$", re.IGNORECASE)


def turkish_lower(text: str) -> str:
    return text.translate(TURKISH_LOWER_MAP).lower()


def build_marker_patterns(markers: list[str]) -> list[tuple[str, re.Pattern]]:
    patterns = []
    for marker in markers:
        normalized = turkish_lower(marker)
        tokens = normalized.split()
        if len(tokens) == 1:
            pattern_str = rf"\b{re.escape(tokens[0])}\w*"
        else:
            parts = [rf"{re.escape(token)}\w*" for token in tokens]
            pattern_str = r"\b" + r"\s+".join(parts)
        patterns.append((normalized, re.compile(pattern_str, re.UNICODE)))
    return patterns


NON_EUROPEAN_MARKER_PATTERNS = build_marker_patterns(NON_EUROPEAN_MARKERS)
DEROGATORY_MARKER_PATTERNS = build_marker_patterns(DEROGATORY_MARKERS)
POST_1453_MARKER_PATTERNS = build_marker_patterns(POST_1453_MARKERS)
MEGARON_NON_GREEK_PATTERNS = build_marker_patterns(MEGARON_NON_GREEK_MARKERS)
SARAZEN_CRUSADER_PATTERNS = build_marker_patterns(SARAZEN_CRUSADER_MARKERS)
SARAZEN_MUSLIM_PATTERNS = build_marker_patterns(SARAZEN_MUSLIM_MARKERS)
SARAZEN_LABELING_PATTERNS = build_marker_patterns(SARAZEN_LABELING_MARKERS)


def ensure_punkt() -> None:
    nltk.download("punkt", quiet=True)
    try:
        nltk.download("punkt_tab", quiet=True)
    except Exception:
        pass


def clean_keyword(keyword: str) -> str:
    cleaned = keyword.strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    cleaned = re.sub(r"[.,;:]+$", "", cleaned)
    return cleaned.strip()


def expand_optional_keyword(keyword: str) -> list[str]:
    match = re.search(r"\(([^)]+)\)", keyword)
    if not match:
        cleaned = clean_keyword(keyword)
        return [cleaned] if cleaned else []
    optional = match.group(1)
    prefix = keyword[: match.start()]
    suffix = keyword[match.end() :]
    base = clean_keyword(prefix + suffix)
    expanded = clean_keyword(prefix + optional + suffix)
    results = []
    if base:
        results.append(base)
    if expanded and expanded != base:
        results.append(expanded)
    return results


def split_keywords_line(line: str) -> list[str]:
    parts = re.split(r"[/,;]", line)
    keywords = []
    for part in parts:
        stripped = part.strip()
        if not stripped:
            continue
        keywords.extend(expand_optional_keyword(stripped))
    return keywords


def load_keywords(path: str) -> list[str]:
    with open(path, "r", encoding="utf-8") as handle:
        lines = handle.read().splitlines()
    keywords = []
    seen = set()
    for line in lines:
        stripped = line.strip()
        if stripped and not stripped.startswith("#"):
            for keyword in split_keywords_line(stripped):
                if not keyword:
                    continue
                normalized = turkish_lower(keyword)
                if normalized in seen:
                    continue
                seen.add(normalized)
                keywords.append(keyword)
    return keywords


def build_root_pattern(normalized_root: str) -> re.Pattern:
    tokens = normalized_root.split()
    if len(tokens) == 1:
        pattern_str = rf"\b{re.escape(tokens[0])}\w*"
    else:
        parts = [rf"{re.escape(token)}\w*" for token in tokens]
        pattern_str = r"\b" + r"\s+".join(parts)
    return re.compile(pattern_str, re.UNICODE)


def build_keyword_patterns(roots: list[str], contextual_map: dict[str, str]) -> list[dict]:
    patterns = []
    for root in roots:
        normalized_root = turkish_lower(root)
        pattern = build_root_pattern(normalized_root)
        patterns.append(
            {
                "root": root,
                "normalized_root": normalized_root,
                "pattern": pattern,
                "context_rule": contextual_map.get(normalized_root, ""),
            }
        )
    return patterns


def normalize_text(text: str) -> str:
    cleaned = text.replace("\n", " ")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def check_ollama_running() -> bool:
    try:
        ollama.list()
        return True
    except Exception:
        return False


def is_roman_numeral(text: str) -> bool:
    cleaned = re.sub(r"[\s\W_]+", "", text, flags=re.UNICODE)
    if not cleaned:
        return False
    return bool(ROMAN_NUMERAL_PATTERN.fullmatch(cleaned))


def categorize_keyword_root(keyword_root: str, context_rule: str) -> str:
    if context_rule:
        return "Bağlamsal Kavramlar"
    if is_roman_numeral(keyword_root):
        return "Roma Rakamları"
    normalized = turkish_lower(keyword_root)
    cleaned = re.sub(r"[\s.]+", "", normalized)
    if cleaned in {"mö", "ms"}:
        return "Gregoryen Miladı"
    if re.search(r"\b(milat|miladi|gregoryan|gregoryen)\b", normalized):
        return "Gregoryen Miladı"
    if re.search(r"\b(çağ|çağı|devri|binyıl|milenyum)\b", normalized):
        return "Çağ ayrımı"
    return "Olgusal Kavramlar"


def build_summary_df(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame(columns=["keyword_root", "count", "category"])

    df["context_rule"] = df["context_rule"].fillna("")
    df["context_rule_match"] = df["context_rule_match"].fillna(False)
    is_contextual = df["context_rule"] != ""

    records = []
    non_context_counts = df.loc[~is_contextual, "keyword_root"].value_counts()
    for keyword_root, count in non_context_counts.items():
        category = categorize_keyword_root(keyword_root, "")
        records.append(
            {"keyword_root": keyword_root, "count": int(count), "category": category}
        )

    context_counts = df.loc[
        is_contextual & (df["context_rule_match"] == True), "keyword_root"
    ].value_counts()
    for keyword_root, count in context_counts.items():
        category = categorize_keyword_root(keyword_root, "context")
        records.append(
            {"keyword_root": keyword_root, "count": int(count), "category": category}
        )

    summary_df = pd.DataFrame(records)
    summary_df["category_order"] = summary_df["category"].map(
        {name: idx for idx, name in enumerate(CATEGORY_ORDER)}
    )
    summary_df = summary_df.sort_values(
        by=["category_order", "count", "keyword_root"],
        ascending=[True, False, True],
    )
    summary_df = summary_df.drop(columns=["category_order"])
    return summary_df[["keyword_root", "count", "category"]]


def build_summary_docx(summary_df: pd.DataFrame, title: str) -> Document:
    document = Document()
    if title:
        document.add_heading(title, level=1)

    if summary_df.empty:
        document.add_paragraph("Veri bulunamadı.")
        return document

    table = document.add_table(rows=1, cols=len(summary_df.columns))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, col_name in enumerate(summary_df.columns):
        header_cells[idx].text = str(col_name)

    for row in summary_df.itertuples(index=False):
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            row_cells[idx].text = "" if pd.isna(value) else str(value)

    return document


def collect_pdf_paths(pdf_arg: str | None, input_dir: Path) -> list[Path]:
    if pdf_arg:
        pdf_path = Path(pdf_arg)
        if pdf_path.is_dir():
            return sorted(pdf_path.glob("*.pdf"))
        if pdf_path.is_file():
            return [pdf_path]
        raise FileNotFoundError(f"PDF path not found: {pdf_path}")

    input_dir.mkdir(parents=True, exist_ok=True)
    return sorted(input_dir.glob("*.pdf"))


def process_pdf(
    pdf_path: Path,
    patterns: list[dict],
    ollama_available: bool,
    model: str,
    output_base_dir: Path,
) -> None:
    rows = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_number, page in enumerate(
            tqdm(pdf.pages, desc=f"Pages ({pdf_path.name})", unit="page"),
            start=1,
        ):
            text = page.extract_text() or ""
            text = normalize_text(text)
            if not text:
                continue

            sentences = sent_tokenize(text, language="turkish")
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                normalized_sentence = turkish_lower(sentence)
                for entry in patterns:
                    for match in entry["pattern"].finditer(normalized_sentence):
                        found_variation = sentence[match.start() : match.end()]
                        rule_match, rule_reason = evaluate_context_rule(
                            entry["context_rule"], normalized_sentence
                        )
                        category = classify_sentence(
                            model=model,
                            word=found_variation,
                            sentence=sentence,
                            ollama_available=ollama_available,
                        )
                        rows.append(
                            {
                                "page_number": page_number,
                                "keyword_root": entry["root"],
                                "found_variation": found_variation,
                                "sentence": sentence,
                                "llm_category": category,
                                "context_rule": entry["context_rule"],
                                "context_rule_match": rule_match,
                                "context_rule_reason": rule_reason,
                            }
                        )

    output_full_dir = output_base_dir / "Çıktılar"
    output_summary_dir = output_base_dir / "özet"
    output_full_dir.mkdir(parents=True, exist_ok=True)
    output_summary_dir.mkdir(parents=True, exist_ok=True)

    output_full = output_full_dir / f"{pdf_path.stem}_çıktı.csv"
    output_summary = output_summary_dir / f"{pdf_path.stem}_özet.docx"

    df = pd.DataFrame(rows)
    df.to_csv(output_full, index=False)
    print(f"Saved {len(df)} rows to {output_full}")

    summary_df = build_summary_df(df)
    summary_doc = build_summary_docx(summary_df, title=f"{pdf_path.stem} Özet")
    summary_doc.save(output_summary)
    print(f"Saved summary to {output_summary}")

def find_first_marker(text: str, patterns: list[tuple[str, re.Pattern]]) -> str:
    for marker, pattern in patterns:
        if pattern.search(text):
            return marker
    return ""


def contains_any_marker(text: str, patterns: list[tuple[str, re.Pattern]]) -> bool:
    return bool(find_first_marker(text, patterns))


def find_year_1453_or_later(text: str) -> str:
    for match in re.finditer(r"\b(\d{4})\b", text):
        year = int(match.group(1))
        if year >= 1453:
            return match.group(1)
    return ""


def is_sarazen_historical_context(text: str) -> bool:
    has_crusader = contains_any_marker(text, SARAZEN_CRUSADER_PATTERNS)
    has_muslim = contains_any_marker(text, SARAZEN_MUSLIM_PATTERNS)
    has_labeling = contains_any_marker(text, SARAZEN_LABELING_PATTERNS)
    return has_crusader and (has_muslim or has_labeling)


def evaluate_context_rule(rule_type: str, normalized_sentence: str) -> tuple[bool, str]:
    if not rule_type:
        return False, ""
    if rule_type == "non_european":
        marker = find_first_marker(normalized_sentence, NON_EUROPEAN_MARKER_PATTERNS)
        return (True, f"non_european_marker:{marker}") if marker else (False, "")
    if rule_type == "post_1453":
        year = find_year_1453_or_later(normalized_sentence)
        if year:
            return True, f"year:{year}"
        marker = find_first_marker(normalized_sentence, POST_1453_MARKER_PATTERNS)
        return (True, f"post_1453_marker:{marker}") if marker else (False, "")
    if rule_type == "derogatory":
        marker = find_first_marker(normalized_sentence, DEROGATORY_MARKER_PATTERNS)
        return (True, f"derogatory_marker:{marker}") if marker else (False, "")
    if rule_type == "megaron_non_greek":
        marker = find_first_marker(normalized_sentence, MEGARON_NON_GREEK_PATTERNS)
        return (True, f"non_greek_marker:{marker}") if marker else (False, "")
    if rule_type == "sarazen":
        if is_sarazen_historical_context(normalized_sentence):
            return False, "historical_usage"
        return True, "non_historical_usage"
    return False, ""


def classify_sentence(
    model: str,
    word: str,
    sentence: str,
    ollama_available: bool,
) -> str:
    if not ollama_available:
        return "LLM_UNAVAILABLE"
    prompt = (
        f"Classify the context of the word '{word}' in this Turkish sentence: "
        f"'{sentence}'. Return one word category."
    )
    try:
        response = ollama.generate(model=model, prompt=prompt)
        raw = str(response.get("response", "")).strip()
        return raw.splitlines()[0] if raw else ""
    except Exception:
        return "LLM_ERROR"


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Find Turkish keyword roots in a PDF and classify context with Ollama."
    )
    parser.add_argument(
        "--pdf",
        default=None,
        help="Path to a PDF file (optional). If omitted, all PDFs in kitaplar are processed.",
    )
    parser.add_argument(
        "--input-dir",
        default="kitaplar",
        help="Folder containing PDFs when --pdf is not provided.",
    )
    parser.add_argument(
        "--keywords",
        default="kelimeler.txt",
        help="Path to kelimeler.txt (one root per line).",
    )
    parser.add_argument(
        "--output-dir",
        "--output",
        dest="output_dir",
        default=".",
        help="Base output directory for Çıktılar/özet folders.",
    )
    parser.add_argument(
        "--model",
        default="llama3.1",
        help="Ollama model name.",
    )
    args = parser.parse_args()

    ensure_punkt()
    output_base_dir = Path(args.output_dir)
    roots = load_keywords(args.keywords)
    if not roots:
        raise ValueError(
            "No keywords loaded. Ensure kelimeler.txt has at least one root."
        )

    contextual_map = {
        turkish_lower(entry["term"]): entry["rule_type"]
        for entry in CONTEXTUAL_RULES
    }
    normalized_roots = {turkish_lower(root) for root in roots}
    for entry in CONTEXTUAL_RULES:
        normalized = turkish_lower(entry["term"])
        if normalized not in normalized_roots:
            roots.append(entry["term"])
            normalized_roots.add(normalized)

    patterns = build_keyword_patterns(roots, contextual_map)
    ollama_available = check_ollama_running()
    if not ollama_available:
        print("Ollama is not reachable; continuing without LLM classification.")

    pdf_paths = collect_pdf_paths(args.pdf, Path(args.input_dir))
    if not pdf_paths:
        print(f"No PDFs found in {Path(args.input_dir)}.")
        return

    for pdf_path in pdf_paths:
        process_pdf(
            pdf_path=pdf_path,
            patterns=patterns,
            ollama_available=ollama_available,
            model=args.model,
            output_base_dir=output_base_dir,
        )


if __name__ == "__main__":
    main()
